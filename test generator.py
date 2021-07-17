import docx
from docx.api import Document
import random

FILENAME = 'תמצית הקורס.docx'
TEST_FILENAME = 'test.docx'


def main():
    doc = Document(docx=FILENAME)
    for table in doc.tables:
        replace_random_paragraphs_from_table(table)
        shuffle_table_rows(table)
    doc.save(TEST_FILENAME)


def replace_random_paragraphs_from_table(table, paragraph_replace_precent=0.5, replace_with='0'):
    """
    Replace random paragraph from table with alternative text.

    :param table:
    :param replace_with:
    :return:
    """
    MAX_NUMBER_TO_RANDOM = 6
    # To not include the header
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if random.randrange(0, MAX_NUMBER_TO_RANDOM) > MAX_NUMBER_TO_RANDOM * paragraph_replace_precent:
                    paragraph.text = replace_with


# Still not working. Invoke error that Table.roww 'can't set attribute'
def shuffle_table_rows(table: docx.Table):
    """
    Shuffle the table rows.
    :param table:
    :return:
    """
    rows_to_shuffle = table.rows[1:]
    random.shuffle(rows_to_shuffle)
    # Be aware to save the first row as header
    table.rows = table.rows[0:1] + rows_to_shuffle


if __name__ == '__main__':
    main()
